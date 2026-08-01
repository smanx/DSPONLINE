from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "教程-DSP极简网络-1.0.15.docx"
QA = ROOT / "artifacts" / "qa"

SECTIONS = [
    ("认识画布", "学会移动、缩放和选择工厂。", "进入任意工厂星球。", ["鼠标左键拖动画布；手机用一根手指拖动。", "滚轮或双指捏合缩放，双击可聚焦画布（如果设置中开启）。", "点击建筑、传送带或空白处查看选择状态。"], "能把工厂移动到屏幕中央，并打开一个建筑检查器。", ["拖动建筑而不是画布。", "手机第二根手指加入后仍继续拖建筑。"], "松开所有手指后点击空白处，再用画布边缘拖动；双指会接管缩放。"),
    ("采集与手搓", "从资源点拿到第一批矿石并制作建筑。", "找到矿脉或原油涌泉，施工托盘中有对应设备。", ["点击资源节点，按住采集或使用采集按钮拿取物品。", "在建造面板打开物品手工制造，选择铁块等中间材料。", "选择建筑并输入数量；材料链会先完整规划，确认后一次扣料。", "把做好的建筑放入光标，再点击画布空白处放置。"], "新建筑出现在画布上，施工托盘数量准确减少。", ["只看到了钢材，却没有铁矿石。", "输入了小数或负数。"], "查看锤子详情中的真实缺料链；原矿不足时回到资源节点采集。"),
    ("放置与堆叠建筑", "让多台相同设备共享一个节点。", "施工托盘中至少有两台相同建筑。", ["选择建筑后点击画布放置第一台。", "选中节点，使用 +1 堆叠或直接输入目标数量。", "需要减少时使用 -1、-10 或输入数量；最低保留一台，完整拆除有单独确认。", "观察输入、输出、燃料、物流槽和缓存随堆叠变化。"], "节点显示新的数量，已有库存与线路物品没有丢失。", ["把货物堆叠误认为建筑数量。", "调低数量时以为超额库存会消失。"], "超额库存会暂时保留并停止新输入；库存降到上限以下后自动恢复。"),
    ("连接传送带", "把生产建筑的输出接到下一个建筑的输入。", "至少有两个兼容端口和对应等级的传送带。", ["从输出端口拖到输入端口；端口吸附范围在手机上更大。", "在线路检查器调整等级、形状、并联数、货物堆叠和优先级。", "普通生产线先获得物资，配送枢纽作为低优先级溢出。", "选中线路后可直接输入并联数量，材料不足会显示缺口。"], "线路变为已连接状态，货物沿线路到达输入缓存。", ["把输出接到输出，或方向反了。", "没有解锁线路等级或施工库存不足。"], "查看端口提示；删除错误线路会返还施工件，已经在途的货物会安全处理。"),
    ("选择与切换配方", "让熔炉和制造台执行正确配方。", "放置熔炉、制造台或其他生产设备。", ["点击建筑后打开选择配方；手机不会自动弹出键盘。", "搜索配方名称或物品名称，点击配方查看耗时和产量。", "确认输入线路与输出空间，再选择配方。", "切换配方前检查现有输入、输出和生产进度。"], "节点显示配方，绿色进度条开始按真实快照推进。", ["移动端直接出现输入法。", "科技未解锁或输入物品不兼容。"], "先关闭配方列表再主动点击搜索框；查看配方详情中的前置科技和缺料原因。"),
    ("物品与物资托盘", "在建筑缓存、手上物品和行星托盘之间安全移动物资。", "身上或建筑里有一件物品。", ["从建筑输入/输出槽点击取出，或拖动物品到光标。", "点击物资托盘图标，把手持物品放回当前行星托盘。", "主动放下手持物品时允许超过托盘单种上限，自动入库仍遵守上限。", "在物资面板可设置托盘上限、删除一半或全部，并进行二次确认。"], "物品数量守恒，光标清空，托盘显示实际数量。", ["把自动输入当成玩家主动放下。", "在错误星球查看托盘。"], "先确认顶部当前行星，再从光标主动放下；失败时不要重复点击，查看提示。"),
    ("打开科技并开始研究", "用矩阵解锁下一阶段设施。", "有科研站和对应颜色矩阵。", ["打开顶部科技入口；再次点击科技按钮可关闭并回到画布。", "选择满足前置条件的科技，确认矩阵输入和研究队列。", "可以暂停、取消或切换研究；已经投入的矩阵和进度会保留。", "无限科技位于同一目录；空结果页可以清除筛选。"], "科技状态变为研究中，进度随模拟时间增加。", ["只看到了当前筛选下的空列表。", "矩阵不足或前置科技未完成。"], "点击清除筛选，打开前置科技；在科技页确认真实缺少哪种矩阵。"),
    ("储物仓、储液罐与配送枢纽", "建立稳定的本地存储和溢出收集。", "已解锁仓储或物资配送枢纽。", ["储物仓和储液罐的输入、输出端口位于卡片边界内；字号变大时会自动换行。", "配送枢纽的三个输入接口可指定物资、恢复自动识别或清空。", "清空有旧线路的接口前先确认断开；缓存和在途货物会安全返还。", "把普通生产线设为高优先级，把配送枢纽留作低优先级溢出。"], "正常产物进入生产线，剩余产物进入当前星球托盘。", ["文字遮住端口。", "误以为低优先级限制最大输送量。"], "调整字体或窗口宽度；检查托盘真实上限，低优先级只改变竞争顺序。"),
    ("行星与星际物流", "让运输机和运输船自动往返。", "完成对应物流科技并放置物流塔。", ["为供给端和需求端配置相同物品槽，设置供需方向和起送比例。", "在塔内输入运输机或运输船目标数量，也可以一键补满。", "跨恒星系航线只由实际派船的一侧检查并预留翘曲器；同一恒星系不消耗。", "观察供电、空闲载具、在途数量和真实阻塞原因。"], "载具从正确的所属站点出发，货物抵达后返回并归还库存。", ["目标端没有翘曲器却误报缺少。", "把执行中的载具直接卸载。"], "按派遣方向检查翘曲器；等待返航后再调整目标数量。"),
    ("翘曲器与跨星系运输", "准备空间翘曲器并理解消耗规则。", "解锁空间翘曲科技和星际物流站。", ["基础配方或重力矩阵绿糖精简配方都可以制造翘曲器。", "把翘曲器装入星际物流站，或打开自动补充并设置目标库存。", "自动补充只读取所在行星托盘；不会跨行星直接取货。", "跨星系每艘往返运输船预留两个翘曲器，取消任务会正确退款。"], "站内库存达到目标，跨星系航线显示可派遣。", ["把同恒星系路线也当成跨星系。", "托盘为空却期待从别的星球补充。"], "切换到物流站所在星球补充托盘，检查路线的星系边界和派遣方向。"),
    ("戴森工程", "从太阳帆和火箭建立持续能源。", "解锁戴森云、弹射器或发射井。", ["在戴森规划中建立轨道和球壳层，记录每层容量。", "电磁轨道弹射器可以指定当前恒星系的目标轨道。", "输入太阳帆、火箭并观察发射、吸附、寿命和接收功率。", "射线接收站分别显示理论接收率、实际利用率和功率瓶颈。"], "轨道与球壳进度增加，接收站输出按 kW 显示。", ["目标轨道被删除或属于其他恒星系。", "接收站输出堵塞却以为是随机效率。"], "重新选择有效轨道，检查输出线路、供电和接收站缓存。"),
    ("建筑制造中心", "自动递归制造复杂建筑并持续补给。", "解锁并放置建筑制造中心，准备原矿和电力。", ["设定建筑目标和数量；高级、稀有和精简配方会优先尝试。", "系统先完整规划多级材料链，再一次性扣除可用原矿和中间材料。", "必要中间材料 WIP 永远保留；额外副产物进托盘，托盘满时记录销毁量。", "检查当前阶段、进度、WIP、缺料、供电和预计时间。"], "铁矿石能自动加工为铁块、钢材再进入目标建筑，不会停在直接材料提示。", ["高级配方未解锁仍强制使用。", "氢等副产物填满缓存后任务停止。"], "系统会回退可完成的基础配方；查看销毁计数和真实阻塞原因，不要手动重启任务。"),
    ("时间扭曲纯挂机", "在不操作工厂的情况下提高真实模拟倍率。", "放置时间扭曲装置并接入足够电力。", ["选中装置并设为主控，选择请求倍率。", "点击“开始纯挂机”，进入独立挂机页面；画布、建造和选择会被冻结。", "挂机页只显示实际倍率、挂机时间、模拟积压、关键产量和保存状态。", "点击“停止挂机”后等待剩余计算和存档校验，再返回工厂。"], "停止后返回同一工厂，暂停期间没有补算，进度从最后有效快照恢复。", ["把挂机页面关闭当成停止。", "供电不足却期待请求倍率完全生效。"], "刷新会恢复最后有效存档；重新进入后查看实际倍率和降档原因，确认后再开始。"),
    ("蓝图与生产区域", "复制布局而不复制矿脉、库存或载具。", "至少有一条线路或一组建筑。", ["框选建筑和线路，保存蓝图并设置名称。", "矿脉只作为资源定位锚点，粘贴时必须落在已有兼容矿脉上。", "部署前确认总需求；库存不足时只部分施工或拒绝，不会复制物资。", "生产区域可拖动四边和四角调整，内部建筑和线路不随区域移动。"], "蓝图部署后线路关系、矿机相对位置和物资守恒保持不变。", ["期待蓝图自动创建矿脉。", "导入区域被蓝图列表遮挡。"], "选择已有资源点再部署；窄窗口请滚动蓝图工作区，导入区与列表独立滚动。"),
    ("常见故障排查", "快速判断为什么产线不动。", "任意生产或物流设备。", ["先看设备状态：缺料、堵塞、供电不足、暂停、缺载具和缺翘曲器是不同原因。", "再看输入输出缓存、线路方向、端口物品和托盘容量。", "打开运营中心 → 性能，按需采样 Worker、传送带、物流和渲染阶段。", "保存前观察主存档校验；失败时使用“立即导出当前进度”。"], "能说出具体阻塞原因并找到对应教程章节。", ["只看到灰色锤子就重复点击。", "把画面刷新慢误认为产量变慢。"], "查看诊断面板和本教程搜索；真实模拟与视觉刷新是分离的，先确认库存和状态数字。"),
    ("性能与存档", "让大型工厂稳定运行并保护进度。", "已进入游戏。", ["设置中的生产画面刷新频率只改变界面发布，不改变产量。", "性能模式只减少粒子、阴影和线路动画；需要时再开启。", "主存档会读回校验，自动快照失败不能阻止主进度保存。", "本地存档、手动槽位和云存档互相独立；导入前先看摘要和校验。"], "切换刷新档位后状态哈希一致，保存失败时界面不会显示假成功。", ["把三秒视觉刷新当成三秒模拟。", "容量不足时反复保存而不导出。"], "先导出当前进度，再管理自动快照；降低视觉效果不会改变存档内容。"),
]

def set_font(run, size=10.5, bold=False, color="283A34"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def label(doc, name, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(f"{name}："), size=10, bold=True, color="1F6B55")
    set_font(p.add_run(text), size=10)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.78)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string("283A34")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2
for style_name, size, color, before, after in [("Heading 1", 17, "1F6B55", 18, 8), ("Heading 2", 13, "267B63", 12, 5)]:
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for style_name in ("List Number", "List Bullet"):
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)
    style.paragraph_format.left_indent = Inches(0.38)
    style.paragraph_format.first_line_indent = Inches(-0.19)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.2
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("DSP极简网络 · 玩家教程 · v1.0.15"), size=8, color="71837B")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("GameState v42 · 教程内容为设备级 UI 资料，不写入存档"), size=8, color="71837B")

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
set_font(title.add_run("DSP极简网络"), size=30, bold=True, color="1F6B55")
title2 = doc.add_paragraph()
title2.paragraph_format.space_after = Pt(4)
set_font(title2.add_run("从零开始的完整玩家教程"), size=22, bold=True, color="142B24")
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
set_font(subtitle.add_run("版本 1.0.15 · 桌面与手机通用 · 面向第一次接触工厂游戏的玩家"), size=11, color="5D756A")
lead = doc.add_paragraph("你不需要玩过《戴森球计划》。这份教程把每个动作拆成可确认的小步骤，并在每节说明成功表现、常见错误和恢复方法。")
for run in lead.runs: set_font(run, size=11)
doc.add_heading("快速开始", level=1)
for text in ["先看第 1-7 节：画布、采集、建筑、线路、配方、托盘和科研。", "再看第 8-12 节：仓储、物流、翘曲器、戴森和建筑制造中心。", "遇到问题时直接搜索“缺料、堵塞、供电、载具、端口或存档”，跳到第 14-15 节。"]:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs: set_font(run)

desktop = QA / "v115-tutorial-desktop.png"
mobile = QA / "v115-tutorial-mobile-390x844.png"
if desktop.exists() or mobile.exists():
    doc.add_heading("界面示意", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for index, path in enumerate((desktop, mobile)):
        cell = table.rows[0].cells[index]
        cell.width = Inches(3.15)
        margins(cell)
        shade(cell, "E8F1EC")
        if path.exists():
            cell.paragraphs[0].add_run().add_picture(str(path), width=Inches(2.95 if index == 0 else 2.25))
        caption = cell.add_paragraph("桌面教程工作区" if index == 0 else "手机教程工作区")
        for run in caption.runs: set_font(run, size=9, bold=True, color="1F6B55")
    p = doc.add_paragraph("游戏内教程会根据屏幕宽度自动切换目录、抽屉和底部操作栏；截图只是布局示意，真实数据仍来自当前工厂。")
    for run in p.runs: set_font(run, size=9, color="5D756A")

doc.add_page_break()
doc.add_heading("目录", level=1)
for index, item in enumerate(SECTIONS, 1):
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run(item[0]), size=10.5, bold=True)

for index, (title_text, goal, prereq, steps, success, mistakes, recovery) in enumerate(SECTIONS, 1):
    doc.add_heading(f"{index}. {title_text}", level=1)
    label(doc, "目标", goal)
    label(doc, "前置条件", prereq)
    doc.add_heading("具体步骤", level=2)
    for step in steps:
        p = doc.add_paragraph(step, style="List Number")
        for run in p.runs: set_font(run)
    doc.add_heading("成功表现", level=2)
    p = doc.add_paragraph(success)
    for run in p.runs: set_font(run, color="1F6B55", bold=True)
    doc.add_heading("常见错误", level=2)
    for mistake in mistakes:
        p = doc.add_paragraph(mistake, style="List Bullet")
        for run in p.runs: set_font(run)
    label(doc, "恢复方法", recovery)
    if index < len(SECTIONS):
        p = doc.add_paragraph("下一节继续：" + SECTIONS[index][0])
        for run in p.runs: set_font(run, size=9, color="71837B")

doc.add_heading("术语表", level=1)
terms = [("物资托盘", "当前行星玩家可直接取放的公共库存。"), ("WIP", "建筑制造中心正在加工、且后续步骤仍需要的中间材料。"), ("并联数", "一条传送带线路同时运行的平行带数量，直接影响吞吐。"), ("起送比例", "物流塔达到该比例后才派遣载具，避免小批量频繁出发。"), ("实际倍率", "供电和设备允许的真实模拟倍率，可能低于请求倍率。")]
table = doc.add_table(rows=1, cols=2)
table.autofit = False
for cell, text in zip(table.rows[0].cells, ("术语", "含义")):
    shade(cell, "E8EEF5")
    margins(cell)
    cell.text = text
    for run in cell.paragraphs[0].runs: set_font(run, size=10, bold=True, color="1F4D78")
for term, definition in terms:
    cells = table.add_row().cells
    for cell in cells: margins(cell)
    cells[0].text, cells[1].text = term, definition
    for run in cells[0].paragraphs[0].runs: set_font(run, size=10, bold=True)
    for run in cells[1].paragraphs[0].runs: set_font(run, size=10)
doc.add_heading("版本与数据边界", level=1)
p = doc.add_paragraph("本教程随客户端版本 1.0.15 发布。教程阅读进度只保存在当前设备的独立 UI 偏好；它不会写入 GameState、手动存档、导入导出文件、云存档或排行榜。玩法数值、配方、缓存、物流和科技状态始终以游戏实际内容目录和模拟引擎为准。")
for run in p.runs: set_font(run)
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
